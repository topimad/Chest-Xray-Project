from django.shortcuts import render, HttpResponse,redirect
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required 
# Create your views here.
import numpy as np
# from tf_keras.saving import load_model
from keras.src.utils import image_utils
import os
from .apps import RadiologyConfig
from django.core.files.storage import FileSystemStorage
from datetime import datetime
from .models import DocumentModel
from docx import Document
from django.conf import settings
from docx.shared import Inches
from .models import DoctorsInfo

# Use the model in your view

def predict_pneumonia(request):
    if request.method == 'POST' and request.FILES.get('image_file'):
        img_file = request.FILES['image_file']
        fs = FileSystemStorage()
        
        filename = fs.save('reports/' + img_file.name, img_file)
        temp_img_path = fs.path(filename)  

        model = RadiologyConfig.model
        img = image_utils.load_img(temp_img_path, target_size=(150, 150))
        img_array = image_utils.img_to_array(img)
        img_array = np.expand_dims(img_array, axis=0) / 255.0
        prediction = model.predict(img_array)
        result = "Pneumonia" if prediction[0][0] > 0.5 else "Normal"

        
        return render(request, 'doctorinterface.html', {'result': result, 'temp_img_path': temp_img_path, 'filename': filename})

    return render(request, 'doctorinterface.html')

#@login_required(login_url='login') # to not get to home without login

def signup(request):
    if request.method=='POST':
        uname=request.POST.get('username')
        email=request.POST.get('email')
        pass1=request.POST.get('password')
        pass2=request.POST.get('password2')
        if pass1!=pass2:
             error_messagx = 'Your password and confirm password are not the same.'
             return render(request,'radiology/signup.html', {'error_messagx': error_messagx})
        my_user=User.objects.create_user(uname,email,pass1)
        my_user.save()
        return redirect('doctor_login')
        #print(uname,email,password,password2) c bon ra tl3o liya f terminal
    return render(request,'radiology/signup.html')


def index(request):
    return render(request,'index.html')


def doctor_login(request):
    if request.method=='POST':
        username=request.POST.get('username')
        pass1=request.POST.get('password') 
        user=authenticate(request,username=username,password=pass1) #the first uname,password from the login the second from the database 
        if user is not None:
            login(request,user)
            return redirect('doctoriterface')
        else :
            error_message = 'userName or Password is incorrect'
            return render(request, 'radiology/login.html', {'error_message': error_message})
        #print(uname,password)
    return render(request,'radiology/login.html')

def LogoutPage(request):
    logout(request)
    return redirect('doctor_login')

def about(request):
    return render(request,'about.html')

def intdoc(request):
    return render(request,'doctorinterface.html')
def blogresult(request):
    return render(request,'blog.html')




def replace_placeholders(doc, placeholders):
    for paragraph in doc.paragraphs:
        if '{doctor_name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{doctor_name}', placeholders.get('doctor_name', ''))
        if '{patient_name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{patient_name}', placeholders.get('patient_name', ''))
        if '{date}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{date}', placeholders.get('date', ''))
        if '{case_description}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{case_description}', placeholders.get('case_description', ''))
        if '{xray_image_placeholder}' in paragraph.text:
            if placeholders.get('xray_image'):
                paragraph.text = ''
                run = paragraph.add_run()
                run.add_picture(placeholders['xray_image'], width=Inches(3))

def generate_report(request):
    if request.method == 'POST':
        firstname = request.POST.get('firstname')
        lastname = request.POST.get('lastname')
        case_description = request.POST.get('case_description')
        temp_img_path = request.POST.get('temp_img_path')

        template_file = os.path.join(settings.MEDIA_ROOT, 'documents', 'Pneumonia_Medical_Report_Template.docx')
        doc = Document(template_file)
        
        user_inputs = {
            'doctor_name': request.user.username,
            'patient_name': f'{firstname} {lastname}',
            'date': datetime.now().strftime('%Y-%m-%d'),
            'case_description': case_description,
            'xray_image': temp_img_path
        }

        replace_placeholders(doc, user_inputs)

        output_file = os.path.join(settings.MEDIA_ROOT, 'documents', f'{firstname}_{lastname}_report.docx')
        doc.save(output_file)

        new_document = DocumentModel(
            user=request.user,
            name=firstname,
            lastname=lastname,
            description=case_description,
            image=temp_img_path,
            file=output_file
        )
        new_document.save()

        # os.system(f'start {output_file}') # c pas la peine wa9ila

        context = {'context': 'Report generated successfully'}
        return render(request, 'doctorinterface.html', context)

    return redirect('doctorinterface')


def types_and_cities(specialiterVilleListe):
    Specialiter = [elm[1] for elm in specialiterVilleListe]
    cities = [elm[0] for elm in specialiterVilleListe]
    return list(set(Specialiter)), list(set(cities))


def dropdownsearch(request):
    DoctorsListe = DoctorsInfo.objects.values_list('City', 'Specialiter').distinct()
    SpecialiterList, citiesList = types_and_cities(DoctorsListe)
    counter = 1
    if request.method == "POST":
        searchSpecialiter = request.POST.get('Specialiter') 
        searchville = request.POST.get('ville')  
        empsearch = DoctorsInfo.objects.filter(Specialiter=searchSpecialiter, City=searchville)
        return render(request, 'guest.html', {"data": empsearch, 'filtered_cities': citiesList, 'filtered_Specialiter': SpecialiterList})
    else:
        displayemp = DoctorsInfo.objects.none()
        return render(request, 'guest.html', {"data": displayemp, 'filtered_cities': citiesList, 'filtered_Specialiter': SpecialiterList , 'counter': counter   })

    
from .models import DocumentModel

@login_required(login_url='doctor_login')
def list_documents(request):
    firstname = request.GET.get('firstname', '')
    lastname = request.GET.get('lastname', '')

    # flwl n affichiw all documents for the user
    documents = DocumentModel.objects.filter(user=request.user)

    if firstname and lastname:
        documents = documents.filter(name__icontains=firstname, lastname__icontains=lastname)
    elif firstname:
        documents = documents.filter(name__icontains=firstname)
    elif lastname:
        documents = documents.filter(lastname__icontains=lastname)
    # if firstname:
    #     documents = documents.filter(name__icontains=firstname)
    # if lastname:
    #     documents = documents.filter(lastname__icontains=lastname)

    return render(request, 'list_documents.html', {'documents': documents})

from django.http import JsonResponse
import json
import pickle
import nltk
from nltk.stem import WordNetLemmatizer
from keras.models import load_model 
import random

# Initialize and load all necessary resources
nltk.download('punkt')
lemmatizer = WordNetLemmatizer()
BASE_DIR = settings.BASE_DIR

model_path = os.path.join(BASE_DIR, 'models', 'MedicalChatbot_model.h5')
data_path = os.path.join(BASE_DIR, 'data', 'intents.json')
words_path = os.path.join(BASE_DIR, 'data', 'words.pkl')
classes_path = os.path.join(BASE_DIR, 'data', 'classes.pkl')

try:
    model = load_model(model_path)
    with open(data_path, 'r') as file:
        intents = json.load(file)
    words = pickle.load(open(words_path, 'rb'))
    classes = pickle.load(open(classes_path, 'rb'))
except Exception as e:
    print(f"Failed to load resources: {e}")

def clean_up_sentence(sentence):
    sentence_words = nltk.word_tokenize(sentence)
    sentence_words = [lemmatizer.lemmatize(word.lower()) for word in sentence_words]
    return sentence_words

def bag_of_words(sentence):
    sentence_words = clean_up_sentence(sentence)
    bag = [0] * len(words)
    for s in sentence_words:
        for i, word in enumerate(words):
            if word == s:
                bag[i] = 1
    return np.array(bag)

def predict_class(sentence):
    bow = bag_of_words(sentence)
    res = model.predict(np.array([bow]))[0]
    ERROR_THRESHOLD = 0.25
    results = [[i, r] for i, r in enumerate(res) if r > ERROR_THRESHOLD]
    results.sort(key=lambda x: x[1], reverse=True)
    return_list = []
    for r in results:
        return_list.append({'intent': classes[r[0]], 'probability': str(r[1])})
    return return_list

def get_response(intents_list, intents_json):
    tag = intents_list[0]['intent']
    list_of_intents = intents_json['intents']
    for i in list_of_intents:
        if i['tag'] == tag:
            return random.choice(i['responses'])
    return "I don't understand your question."

def chatbot_view(request):
    user_input = request.GET.get('msg', '')
    if user_input:
        try:
            prediction = predict_class(user_input)
            response = get_response(prediction, intents)
            return JsonResponse({'response': response})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return JsonResponse({'response': 'No input provided'})